"""Document text extraction for the optional local knowledge backend."""

from __future__ import annotations

import mimetypes
from pathlib import Path
from typing import Dict, List

from .models import DependencyStatus, DocumentPage, ExtractedDocument


SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt", ".md", ".markdown"}


class ExtractionError(Exception):
    """Raised when a document cannot be extracted."""


class UnsupportedDocumentError(ExtractionError):
    """Raised when a file suffix is outside the supported upload set."""


class MissingDependencyError(ExtractionError):
    """Raised when an optional parser dependency is not installed."""


def dependency_status() -> Dict[str, DependencyStatus]:
    """Return parser dependency status without importing the whole backend."""

    return {
        "sqlite3": _dependency("sqlite3", "SQLite metadata storage"),
        "pypdf": _dependency("pypdf", "PDF text extraction"),
        "python-docx": _dependency("docx", "DOCX text extraction"),
    }


def extract_document(path: Path) -> ExtractedDocument:
    """Extract text from a supported document file."""

    path = Path(path)
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise UnsupportedDocumentError(
            f"unsupported document type '{suffix}'. Supported: {', '.join(sorted(SUPPORTED_SUFFIXES))}"
        )

    if suffix == ".pdf":
        pages = _extract_pdf(path)
    elif suffix == ".docx":
        pages = _extract_docx(path)
    else:
        pages = _extract_text(path)

    mime_type = mimetypes.guess_type(path.name)[0] or "application/octet-stream"
    return ExtractedDocument(
        title=path.stem,
        source_path=str(path),
        mime_type=mime_type,
        pages=pages,
        metadata={"suffix": suffix, "page_count": len(pages)},
    )


def _extract_pdf(path: Path) -> List[DocumentPage]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise MissingDependencyError("pypdf is required for PDF ingestion") from exc

    reader = PdfReader(str(path))
    pages: List[DocumentPage] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(DocumentPage(page=index, text=text.strip()))
    if not pages:
        pages.append(DocumentPage(page=1, text=""))
    return pages


def _extract_docx(path: Path) -> List[DocumentPage]:
    try:
        from docx import Document
    except ImportError as exc:
        raise MissingDependencyError("python-docx is required for DOCX ingestion") from exc

    doc = Document(str(path))
    parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return [DocumentPage(page=1, text="\n".join(parts))]


def _extract_text(path: Path) -> List[DocumentPage]:
    for encoding in ("utf-8-sig", "utf-8", "gbk", "gb2312", "latin-1"):
        try:
            return [DocumentPage(page=1, text=path.read_text(encoding=encoding))]
        except UnicodeError:
            continue
    raise ExtractionError("unable to decode text document")


def _dependency(module_name: str, detail: str) -> DependencyStatus:
    try:
        __import__(module_name)
        return DependencyStatus(name=module_name, available=True, detail=detail)
    except Exception as exc:
        return DependencyStatus(name=module_name, available=False, detail=str(exc) or detail)
